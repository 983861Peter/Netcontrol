from docx import Document
from docx.shared import Pt

doc = Document()

doc.add_heading('System Code Documentation', level=0)

doc.add_paragraph('Project: ACS 400 PROJECT')
doc.add_paragraph('Generated documentation for the router simulation backend system code. This document summarizes the core modules, data models, API behavior, device communication interfaces, and operational flow.')

doc.add_heading('1. System Overview', level=1)
doc.add_paragraph('The system is a router management and simulation platform built around a FastAPI backend in the `router_sim` package. It provides authentication, device management, router configuration deployment, event logging, background monitoring, and WebSocket alert broadcasting. Static web UI assets are served from `router_sim/ui`.')

doc.add_heading('2. Core Components', level=1)

doc.add_paragraph('The primary backend modules are:')
components = [
    ('router_sim/router_api.py', 'Main FastAPI application entrypoint. Sets up CORS, static UI mounting, default administrator initialization, database session management, and helper methods for adapter loading, event logging, and config deployment.'),
    ('router_sim/models.py', 'SQLAlchemy ORM definitions for the core data model, including AuthUser, Device, Company, TransmissionStation, Sector, CompanyDefaults, DeviceTypeTemplate, NetworkInterface, and related entities.'),
    ('router_sim/routes.py', 'Authentication and alert routing logic: login/logout, current user retrieval, role checks, event logging, and WebSocket alert route handling.'),
    ('router_sim/admin_routes.py', 'Admin-level HTTP routes for user and system management, including user registration, user management, and elevated admin actions.'),
    ('router_sim/dlink_adapter.py', 'Vendor adapter for D-Link routers. Detects protocol availability (Telnet / HTTP) and supports login, status retrieval, WiFi configuration updates, and reboot operations.'),
    ('router_sim/device_comm.py', 'Basic router communication utilities for ping and SSH command execution. Provides a generic device interface that can be extended for additional brands.'),
    ('router_sim/poller.py', 'Background monitoring loop that polls active devices, updates state, evaluates resets, and triggers automatic restore workflows when resets are detected.'),
    ('router_sim/ws_broadcast.py', 'WebSocket client registry and broadcast helper used to deliver real-time event notifications to connected UI clients.'),
]
for name, desc in components:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(name + ': ').bold = True
    p.add_run(desc)

doc.add_heading('3. API Summary', level=1)
doc.add_paragraph('The backend uses FastAPI and defines both application-level routes and sub-routers for authentication, backup, and device operations. Key endpoint groups include:')
api_points = [
    ('Authentication', 'Login, logout, user context retrieval, and token validation. Uses JWT Bearer tokens and protects endpoint access by user role.'),
    ('Router management', 'Device discovery, configuration deployment, backup creation, router status updates, and device reset handling.'),
    ('WebSocket alerts', 'Real-time alert broadcast over /ws/alerts and event notification support for frontend clients.'),
    ('Static UI assets', 'Serves the `router_sim/ui` static web application under /static if the directory exists.'),
]
for title, desc in api_points:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title + ': ').bold = True
    p.add_run(desc)

doc.add_heading('4. Data Model Summary', level=1)
doc.add_paragraph('Key persistent objects represented in the database include:')
models = [
    ('AuthUser', 'System users with authentication credentials, role, activity flags, and company association.'),
    ('Device', 'Network devices managed by the system. Stores identifiers, IP configuration, credentials, device type, metadata, and restore state.'),
    ('Company', 'Tenant or operator entity owning devices, users, clients, and configuration templates.'),
    ('TransmissionStation', 'Logical network station or gateway point for wireless service providers.'),
    ('Sector', 'A sector or cell served by a transmission station, including IP assignment mode and device model details.'),
    ('CompanyDefaults', 'Default network configuration values for a company, such as DNS, NTP, syslog, and management credentials.'),
    ('DeviceTypeTemplate', 'Predefined device configuration templates that can be applied across devices.'),
    ('NetworkInterface', 'Interface-level details for managed devices, including IP addresses, operational state, counters, and VLAN or neighbor data.'),
]
for name, desc in models:
    p = doc.add_paragraph(style='List Number')
    p.add_run(name + ': ').bold = True
    p.add_run(desc)

doc.add_heading('5. Device Communication and Adapter Architecture', level=1)
doc.add_paragraph('The system separates device interaction from backend orchestration:')
communication = [
    ('Generic device comm', 'router_sim/device_comm.py provides basic ping and SSH command execution capabilities for generic router interaction.'),
    ('Vendor adapters', 'router_sim/dlink_adapter.py is an example vendor adapter. It dynamically selects Telnet or HTTP protocol, performs login, gathers status, updates WiFi settings, and reboots devices.'),
    ('Dynamic adapter loading', 'router_api.py uses adapter_loader.get_vendor_adapter to load vendor-specific adapter modules based on device model or vendor name.'),
]
for title, desc in communication:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title + ': ').bold = True
    p.add_run(desc)

doc.add_heading('6. Monitoring and Auto-Restore Behavior', level=1)
doc.add_paragraph('The backend includes a continuous monitoring engine that evaluates device health and restore needs:')
monitoring = [
    ('Polling loop', 'router_sim/poller.py polls devices every 10 seconds, updates device state, and uses configured adapters to fetch runtime metrics.'),
    ('Reset detection', 'After polling, devices are analyzed for factory resets or configuration loss, and a restore flag is set if reset conditions are met.'),
    ('Auto restore', 'If a recent backup exists, the system attempts to restore device configuration automatically using the appropriate adapter.'),
    ('Event logging', 'Reset and restore events are persisted to EventLog and broadcast to connected WebSocket clients.'),
]
for title, desc in monitoring:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title + ': ').bold = True
    p.add_run(desc)

doc.add_heading('7. Startup and Initialization', level=1)
doc.add_paragraph('Important application startup behavior:')
startup_points = [
    ('Background tasks', 'router_api.py schedules `monitoring_loop()` on startup so device polling begins when the app launches.'),
    ('Default admin creation', 'A default administrator account is created on startup if no admin user exists, using environment variables for username, password, and email.'),
    ('Database setup', 'The app relies on SQLAlchemy sessions from router_sim/db.py and model metadata for table creation, though table creation appears commented out in startup hooks.'),
]
for title, desc in startup_points:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title + ': ').bold = True
    p.add_run(desc)

doc.add_heading('8. Frontend and User Interface', level=1)
doc.add_paragraph('UI content is stored in `router_sim/ui`. The backend serves the static files at `/static` if the directory exists. This means the browser app is bundled separately from the API and is published through FastAPI static file routing.')

doc.add_heading('9. Notes and Recommendations', level=1)
doc.add_paragraph('Potential improvement areas in the codebase:')
notes = [
    ('Database migrations', 'Add explicit migration support (Alembic or similar) rather than relying on SQLAlchemy table creation comments.'),
    ('Error handling', 'Expand exception handling in adapter and background tasks to avoid loss of async failures.'),
    ('Schema documentation', 'Document the public API routes and payload contracts more explicitly with OpenAPI schemas or a README.'),
    ('Adapter extensibility', 'Add a common adapter base class to standardize vendor adapters and reduce protocol-specific branching.'),
]
for title, desc in notes:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title + ': ').bold = True
    p.add_run(desc)

doc.save('System_Code_Documentation.docx')
print('System_Code_Documentation.docx generated successfully.')
